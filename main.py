import streamlit as st
import pandas as pd
import pdf2docx
import fitz
import tempfile
from docx import Document
import docx2pdf
import pdfplumber
from reportlab.lib.pagesizes import letter
from docx import Document
from docx2pdf import convert
import os
from reportlab.pdfgen import canvas
from fpdf import FPDF

def converter_pdf_excel(upload_arquivo_pdf):
    tables = []
    with pdfplumber.open(upload_arquivo_pdf) as pdf:
        for page in pdf.pages:
            table = page.extract_table()
            if table:
                tables.append(table)
    if not tables:
        st.error("Nenhuma tabela encontrada no PDF.")
        return None

    with pd.ExcelWriter('output.xlsx') as writer:
        for i, table in enumerate(tables):
            df = pd.DataFrame(table[1:], columns=table[0])
            df.to_excel(writer, sheet_name=f'Sheet_{i+1}', index=False)
    return 'output.xlsx'

def converter_pdf_docx(upload_arquivo_pdf):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        temp_pdf.write(upload_arquivo_pdf.read())
        temp_pdf_path = temp_pdf.name
    cv = pdf2docx.Converter(temp_pdf_path)
    cv.convert('output.docx')
    return 'output.docx'

def converter_pdf_para_pdfa(upload_arquivo_pdf):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(upload_arquivo_pdf.read())
            temp_pdf_path = temp_pdf.name
        input_pdf = fitz.open(temp_pdf_path)
        input_pdf.set_metadata({
            "producer": "PyMuPDF",
            "creator": "PyMuPDF",
            "author": "Author Name",
            "subject": "Subject",
            "title": "PDF/A Document",
            "format": "PDF/A",
        })
        output_pdf_path = 'output.pdfa'
        input_pdf.save(output_pdf_path)
        return output_pdf_path
    except Exception as e:
        st.error(f"Erro na conversão para PDFA: {e}")
        return None

def converter_docx_para_pdf(upload_arquivo_docx):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(upload_arquivo_docx.read())
        temp_docx_path = temp_docx.name
        output_pdf_path = 'output.pdf'
        
        doc = fitz.open(temp_docx_path)
        
        doc.save(output_pdf_path)
        
        doc.close()
        
        return output_pdf_path

def converter_docx_para_xlsx(upload_arquivo_docx):
    document = Document(upload_arquivo_docx)
    data = [[cell.text for cell in row.cells] for row in document.tables[0].rows]
    df = pd.DataFrame(data)
    df.to_excel('output.xlsx', index=False)

    return 'output.xlsx'

def converter_docx_para_pdfa(upload_arquivo_docx):
    pdf_path = converter_docx_para_pdf(upload_arquivo_docx)
    try:
        input_pdf = fitz.open(pdf_path)
        input_pdf.set_metadata({
            "producer": "PyMuPDF",
            "creator": "PyMuPDF",
            "author": "Author Name",
            "subject": "Subject",
            "title": "PDF/A Document",
            "format": "PDF/A",
        })
        output_pdf_path = 'output.pdfa'
        input_pdf.save(output_pdf_path)
        return output_pdf_path
    except Exception as e:
        st.error(f"Erro na conversão para PDFA: {e}")
        return None

def converter_xlsx_para_pdf(upload_arquivo_xlsx):
    df = pd.read_excel(upload_arquivo_xlsx)
    output_pdf_path = 'output.pdf'

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size = 12)

    for col in df.columns:
        pdf.cell(40, 10, col, border=1, align='C')
    pdf.ln(10)

    for row in df.values:
        for value in row:
            pdf.cell(40, 10, str(value), border=1, align='C')
        pdf.ln(10)

    pdf.output(output_pdf_path)

    return output_pdf_path

def converter_xlsx_para_docx(upload_arquivo_xlsx):
    df = pd.read_excel(upload_arquivo_xlsx)
    doc = Document()
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    doc.save('output.docx')

    return 'output.docx'

def converter_xlsx_para_pdfa(upload_arquivo_xlsx):
    pdf_path = converter_xlsx_para_pdf(upload_arquivo_xlsx)
    try:
        input_pdf = fitz.open(pdf_path)
        input_pdf.set_metadata({
            "producer": "PyMuPDF",
            "creator": "PyMuPDF",
            "author": "Author Name",
            "subject": "Subject",
            "title": "PDF/A Document",
            "format": "PDF/A",
        })
        output_pdf_path = 'output.pdfa'
        input_pdf.save(output_pdf_path)
        return output_pdf_path
    except Exception as e:
        st.error(f"Erro na conversão para PDFA: {e}")
        return None

def ui_up_pdf_xlsx():
    up_pdf_para_xlsx = st.file_uploader("📎 Conversor de PDF para XLSX", type="pdf", key="pdf_to_xlsx")
    if up_pdf_para_xlsx is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_pdf_excel(up_pdf_para_xlsx)
        if arquivo_saida:
            barra_progresso.progress(100)
            st.success("Arquivo convertido com sucesso! 🎉")
            with open(arquivo_saida, 'rb') as f:
                st.download_button('Download 💾', f, file_name='output.xlsx')

def ui_up_pdf_docx():
    up_pdf_para_docx = st.file_uploader("📎 Conversor de PDF para DOCX", type="pdf", key="pdf_to_docx")
    if up_pdf_para_docx is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_pdf_docx(up_pdf_para_docx)
        barra_progresso.progress(100)
        st.success("Arquivo convertido com sucesso! 🎉")
        with open(arquivo_saida, 'rb') as f:
            st.download_button('Download 💾', f, file_name='output.docx')

def ui_up_pdf_pdfa():
    up_pdf_para_pdfa = st.file_uploader("📎 Conversor de PDF para PDFA", type="pdf", key="pdf_to_pdfa")
    if up_pdf_para_pdfa is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_pdf_para_pdfa(up_pdf_para_pdfa)
        if arquivo_saida:
            barra_progresso.progress(100)
            st.success("Arquivo convertido com sucesso! 🎉")
            with open(arquivo_saida, 'rb') as f:
                st.download_button('Download 💾', f, file_name='output.pdfa')

def ui_up_docx_pdf():
    up_docx_para_pdf = st.file_uploader("📎 Conversor de DOCX para PDF", type="docx", key="docx_to_pdf")
    if up_docx_para_pdf is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_docx_para_pdf(up_docx_para_pdf)
        barra_progresso.progress(100)
        st.success("Arquivo convertido com sucesso! 🎉")
        with open(arquivo_saida, 'rb') as f:
            st.download_button('Download 💾', f, file_name='output.pdf')

def ui_up_docx_xslx():
    up_docx_para_xslx = st.file_uploader("📎 Conversor de DOCX para XLSX", type="docx", key="docx_to_xlsx")
    if up_docx_para_xslx is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_docx_para_xlsx(up_docx_para_xslx)
        barra_progresso.progress(100)
        st.success("Arquivo convertido com sucesso! 🎉")
        with open(arquivo_saida, 'rb') as f:
            st.download_button('Download 💾', f, file_name='output.xlsx')

def ui_up_docx_pdfa():
    up_docx_para_pdfa = st.file_uploader("📎 Conversor de DOCX para PDFA", type="docx", key="docx_to_pdfa")
    if up_docx_para_pdfa is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_docx_para_pdfa(up_docx_para_pdfa)
        if arquivo_saida:
            barra_progresso.progress(100)
            st.success("Arquivo convertido com sucesso! 🎉")
            with open(arquivo_saida, 'rb') as f:
                st.download_button('Download 💾', f, file_name='output.pdf')

def ui_up_xslx_pdf():
    up_xslx_pdf = st.file_uploader("📎 Conversor de XSLX para PDF", type="xlsx", key="xlsx_to_pdf")
    if up_xslx_pdf is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_xlsx_para_pdf(up_xslx_pdf)
        barra_progresso.progress(100)
        st.success("Arquivo convertido com sucesso! 🎉")
        with open(arquivo_saida, 'rb') as f:
            st.download_button('Download 💾', f, file_name='output.pdf')

def ui_up_xslx_docx():
    up_xslx_docx = st.file_uploader("📎 Conversor de XSLX para DOCX", type="xlsx", key="xlsx_to_docx")
    if up_xslx_docx is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_xlsx_para_docx(up_xslx_docx)
        barra_progresso.progress(100)
        st.success("Arquivo convertido com sucesso! 🎉")
        with open(arquivo_saida, 'rb') as f:
            st.download_button('Download 💾', f, file_name='output.docx')

def ui_up_xslx_pdfa():
    up_xlsx_pdfa = st.file_uploader("📎 Conversor de XSLX para PDFA", type="xlsx", key="xlsx_to_pdfa")
    if up_xlsx_pdfa is not None:
        barra_progresso = st.progress(0)
        arquivo_saida = converter_xlsx_para_pdfa(up_xlsx_pdfa)
        if arquivo_saida:
            barra_progresso.progress(100)
            st.success("Arquivo convertido com sucesso! 🎉")
            with open(arquivo_saida, 'rb') as f:
                st.download_button('Download 💾', f, file_name='output.pdfa')

def update_access_counter():
    try:
        with open('access_counter.txt', 'r') as f:
            count = int(f.read())
    except FileNotFoundError:
        count = 0

    count += 1

    with open('access_counter.txt', 'w') as f:
        f.write(str(count))

    return count


def page_conversor():
    st.header('📂 CONVERSOR DE ARQUIVOS', divider=True)

    tab1, tab2, tab3 = st.tabs(['PDF', 'DOCX', 'XLSX'])

    with tab1:
        ui_up_pdf_xlsx()
        ui_up_pdf_docx()
        ui_up_pdf_pdfa()

    with tab2:
        ui_up_docx_pdf()
        ui_up_docx_xslx()
        ui_up_docx_pdfa()

    with tab3:
        ui_up_xslx_pdf()
        ui_up_xslx_docx()
        ui_up_xslx_pdfa()

def sidebar():
    st.sidebar.header("Creditos", divider=True)
    st.sidebar.markdown("""
    💜 Gostou do projeto? Me pague um café!
    
    **PIX:** hugorogerio522@gmail.com""")


def main():
    st.set_page_config(page_title="Conversor de arquivos", page_icon="📂")
    sidebar()

    access_count = update_access_counter()
    st.sidebar.markdown(f"👥 Número de acessos: {access_count}")
    
    page_conversor()

if __name__ == "__main__":
    main()
